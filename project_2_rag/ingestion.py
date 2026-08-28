import csv, io, re
from pathlib import Path

def extract_text(name,data):
    ext=Path(name).suffix.lower()
    if ext in {".md",".txt"}: return data.decode("utf-8",errors="replace")
    if ext==".csv":
        rows=list(csv.reader(io.StringIO(data.decode("utf-8",errors="replace"))))
        return "\n".join(" | ".join(row) for row in rows)
    if ext==".pdf":
        from pypdf import PdfReader
        return "\n\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(data)).pages)
    if ext==".docx":
        from docx import Document
        d=Document(io.BytesIO(data));parts=[]
        for para in d.paragraphs:
            if para.text.strip():parts.append(para.text.strip())
        for table in d.tables:
            for row in table.rows:parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, Markdown, TXT, or CSV.")

def heading_sections(text):
    lines=text.splitlines();sections=[];heading="Document content";body=[]
    def flush():
        nonlocal body
        content="\n".join(body).strip()
        if content:sections.append((heading,content))
        body=[]
    for line in lines:
        s=line.strip()
        markdown=re.match(r"^#{1,6}\s+(.+)$",s)
        numbered=re.match(r"^(?:\d+(?:\.\d+)*[.)]?\s+)(.{3,100})$",s)
        title_case=bool(s and len(s)<90 and s==s.title() and not s.endswith('.'))
        if markdown or numbered or title_case:
            flush();heading=(markdown.group(1) if markdown else numbered.group(1) if numbered else s).strip()
        else:body.append(line)
    flush()
    return sections or [("Document content",text.strip())]

def chunk_document(text,metadata,max_words=180,overlap=35):
    chunks=[];cid=0
    for section,body in heading_sections(text):
        words=body.split()
        step=max(1,max_words-overlap)
        for start in range(0,len(words),step):
            piece=" ".join(words[start:start+max_words]).strip()
            if not piece:continue
            cid+=1
            chunks.append({**metadata,"section":section,"chunk":cid,"text":piece,"word_count":len(piece.split()),"source_id":f"S{cid}"})
            if start+max_words>=len(words):break
    return chunks

def builtin_documents(directory):
    docs=[]
    defaults={"torque_control.md":{"title":"Torque Control Guidance","document_number":"WI-FAST-014","revision":"C","plant":"All","process":"Fastening","document_type":"Work Instruction","status":"Released"},
              "nonconformance.md":{"title":"Nonconforming Product Control","document_number":"QP-NCR-002","revision":"B","plant":"All","process":"Quality","document_type":"Procedure","status":"Released"},
              "pfmea.md":{"title":"PFMEA Review Guidance","document_number":"QP-PFMEA-001","revision":"A","plant":"All","process":"Quality Planning","document_type":"Procedure","status":"Released"}}
    for p in sorted(Path(directory).glob("*.md")):
        meta={"source":p.name,"effective_date":"2026-01-01","owner":"Quality Systems","confidentiality":"Internal non-sensitive",**defaults.get(p.name,{"title":p.stem.replace('_',' ').title(),"document_number":p.stem.upper(),"revision":"A","plant":"All","process":"General","document_type":"Guidance","status":"Released"})}
        docs.extend(chunk_document(p.read_text(encoding="utf-8"),meta))
    return docs
